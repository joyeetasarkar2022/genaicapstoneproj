import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd
import yaml
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

from app.core.security import sanitize_text

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    """Extract raw text from supported source formats."""

    def load(self, path: Path) -> list[LoadedDocument]:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == '.pdf':
            return self._load_pdf(path)
        if suffix in {'.txt', '.md'}:
            return [self._document(path.read_text(encoding='utf-8', errors='ignore'), path)]
        if suffix == '.csv':
            return [self._load_csv(path)]
        if suffix in {'.xlsx', '.xls'}:
            return self._load_excel(path)
        if suffix == '.json':
            return [self._load_json(path)]
        if suffix in {'.yaml', '.yml'}:
            return [self._load_yaml(path)]
        if suffix == '.docx':
            return [self._load_docx(path)]
        raise ValueError(f'Unsupported file type: {suffix}')

    def _document(self, text: str, path: Path, **extra: Any) -> LoadedDocument:
        text = sanitize_text(text)
        metadata = {
            'source': path.name,
            'file_path': str(path),
            'extension': path.suffix.lower(),
            **extra,
        }
        return LoadedDocument(text=text, metadata=metadata)

    def _load_pdf(self, path: Path) -> list[LoadedDocument]:
        reader = PdfReader(str(path))
        docs: list[LoadedDocument] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ''
            except Exception as exc:  # pragma: no cover - depends on PDF encoding
                logger.warning('Failed to extract page %s from %s: %s', index, path, exc)
                text = ''
            if sanitize_text(text):
                docs.append(self._document(text, path, page=index))
        return docs

    def _load_csv(self, path: Path) -> LoadedDocument:
        frame = pd.read_csv(path)
        text = frame.to_csv(index=False)
        return self._document(text, path, rows=int(frame.shape[0]), columns=int(frame.shape[1]))

    def _load_excel(self, path: Path) -> list[LoadedDocument]:
        sheets = pd.read_excel(path, sheet_name=None)
        docs: list[LoadedDocument] = []
        for sheet_name, frame in sheets.items():
            text = f'Sheet: {sheet_name}\n' + frame.to_csv(index=False)
            docs.append(self._document(text, path, sheet=sheet_name, rows=int(frame.shape[0]), columns=int(frame.shape[1])))
        return docs

    def _load_json(self, path: Path) -> LoadedDocument:
        data = json.loads(path.read_text(encoding='utf-8', errors='ignore'))
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return self._document(text, path)

    def _load_yaml(self, path: Path) -> LoadedDocument:
        data = yaml.safe_load(path.read_text(encoding='utf-8', errors='ignore'))
        text = yaml.safe_dump(data, sort_keys=False, allow_unicode=True)
        return self._document(text, path)

    def _load_docx(self, path: Path) -> LoadedDocument:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return self._document('\n'.join(paragraphs), path)
